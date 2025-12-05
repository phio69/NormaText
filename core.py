"""
Ядро логики NormaText: работа с документом, проверка по ГОСТ, лемматизация.
"""

from docx import Document
import pymorphy3
from dataclasses import dataclass
from typing import List, Tuple, Optional
import re
from dictionaries import FORBIDDEN_WORDS, TERMINOLOGY_REPLACEMENTS

@dataclass
class Heading:
    # Класс для представления заголовка
    level: int           # Уровень (1, 2, 3...)
    text: str           # Текст заголовка
    paragraph_index: int # Номер абзаца в документе
    number: str         # Номер (например, "1", "1.1", "2.3.1")

# Инициализация морфологического анализатора (один раз для всего приложения)
_morph = pymorphy3.MorphAnalyzer()

def load_document(file_path):
    # Загружает .docx-документ.
    return Document(file_path)

def save_fixed_document(document, original_path):
    # Сохраняет исправленный документ
    from pathlib import Path

    original = Path(original_path)
    new_name = f"{original.stem}_исправленный{original.suffix}"
    new_path = original.parent / new_name

    try:
        document.save(new_path)
        return f"Документ сохранён: {new_path}"
    except Exception as e:
        return f"Ошибка сохранения: {str(e)}"

def check_terminology(document):
    # Проверяет документ на наличие запрещённых слов.
    errors = []
    all_lines = [(i + 1, p.text) for i, p in enumerate(document.paragraphs) if p.text.strip()]

    for line_num, text in all_lines:
        words = text.split()
        for word in words:
            clean = word.strip(".,;:!?\"'()[]{}—–-")
            if not clean or not clean.isalpha():
                continue
            try:
                normal_form = _morph.parse(clean.lower())[0].normal_form
                if normal_form in FORBIDDEN_WORDS:
                    errors.append(
                        f"• Стр. {line_num}: Недопустимое слово «{clean}» (основа: «{normal_form}»)"
                    )
            except Exception:
                continue
    return errors


def auto_fix_terminology(document):
    # Автоматически заменяет запрещенные слова на корректные аналоги
    # Возвращает количество выполненных замен
    replacements_count = 0

    for paragraph in document.paragraphs:
        if not paragraph.text.strip():
            continue

        # Разбиваем абзац на слова и обрабатываем каждое
        original_text = paragraph.text
        words = original_text.split()
        new_words = []

        for word in words:
            clean_word = word.strip(".,;:!?\"'()[]{}—–-")

            if not clean_word or not clean_word.isalpha():
                new_words.append(word)
                continue

            try:
                normal_form = _morph.parse(clean_word.lower())[0].normal_form

                if normal_form in TERMINOLOGY_REPLACEMENTS:
                    replacement = TERMINOLOGY_REPLACEMENTS[normal_form]

                    if replacement:  # Если замена не пустая
                        # Сохраняем оригинальное форматирование (регистр)
                        if clean_word.istitle():
                            replacement = replacement.title()
                        elif clean_word.isupper():
                            replacement = replacement.upper()

                        # Заменяем слово в оригинальном тексте с сохранением знаков препинания
                        new_word = word.replace(clean_word, replacement)
                        new_words.append(new_word)
                        replacements_count += 1
                    else:
                        # Если замена пустая - удаляем слово
                        replacements_count += 1
                        continue
                else:
                    new_words.append(word)

            except Exception:
                new_words.append(word)

        # Обновляем текст абзаца
        if replacements_count > 0:
            paragraph.text = ' '.join(new_words)

    return replacements_count

def check_structure(document, doc_type: str) -> List[str]:
    # Проверяет структуру документа по ГОСТу
    errors = []

    # Используем новую функцию проверки реквизитов
    errors.extend(check_required_fields(document, doc_type))

    # Дополнительные проверки структуры
    text_sample = " ".join(p.text for p in document.paragraphs[:5])
    #Проверка Оформления
    errors.extend(check_formatting(document))
    errors.extend(check_paragraphs_structure(document))
    errors.extend(check_lists_formatting(document))
    errors.extend(check_fonts_and_sizes(document))

    # Проверка наличия даты
    if "202" not in text_sample:
        errors.append("• Возможно отсутствует дата документа")

    return errors


def check_fonts_and_sizes(document) -> List[str]:
    # Проверка шрифтов и размеров по ГОСТу
    errors = []
    non_times_fonts = set()
    wrong_sizes = set()

    for i, paragraph in enumerate(document.paragraphs):
        if not paragraph.text.strip():
            continue

        is_heading = _is_heading(paragraph)

        for run in paragraph.runs:
            if not run.text.strip():
                continue

            # Проверка шрифта
            if hasattr(run, 'font') and run.font.name:
                font_name = run.font.name.lower()
                allowed_fonts = ['times', 'times new roman']

                if not any(allowed in font_name for allowed in allowed_fonts):
                    non_times_fonts.add(run.font.name)

            # Проверка размера шрифта
            if hasattr(run, 'font') and hasattr(run.font, 'size') and run.font.size:
                try:
                    # Конвертируем в пункты (1 point = 12700)
                    if hasattr(run.font.size, 'pt'):
                        size_pt = run.font.size.pt
                    else:
                        size_pt = run.font.size / 12700

                    if is_heading:
                        # Для заголовков: 14-16 pt
                        if size_pt < 14 or size_pt > 16:
                            wrong_sizes.add(f"заголовок: {size_pt:.1f}pt (требуется 14-16pt)")
                    else:
                        # Для основного текста: 12-14 pt
                        if size_pt < 12 or size_pt > 14:
                            wrong_sizes.add(f"текст: {size_pt:.1f}pt (требуется 12-14pt)")

                except (AttributeError, TypeError):
                    continue

    # Формируем ошибки
    if non_times_fonts:
        fonts_list = ', '.join(non_times_fonts)
        errors.append(f"• Обнаружены нерекомендуемые шрифты: {fonts_list} (ГОСТ: Times New Roman)")

    if wrong_sizes:
        sizes_list = '; '.join(wrong_sizes)
        errors.append(f"• Несоответствие размеров шрифта: {sizes_list}")

    return errors

def check_formatting(document) -> List[str]:
    # Проверяет базовое оформление документа по ГОСТ Р 7.0.97-2016
    errors = []

    for i, paragraph in enumerate(document.paragraphs):
        if not paragraph.text.strip():
            continue

        # Проверка выравнивания (ГОСТ: по ширине для основного текста)
        if hasattr(paragraph, 'alignment') and paragraph.alignment:
            # 0=left, 1=center, 2=right, 3=justify
            if paragraph.alignment not in [0, 3]:  # Допустимо: по левому краю и по ширине
                errors.append(f"• Стр. {i + 1}: Рекомендуется выравнивание по ширине или левому краю")

        # Проверка на использование CAPSLOCK (не рекомендуется)
        text = paragraph.text
        if len(text) > 10 and text.isupper():
            errors.append(f"• Стр. {i + 1}: Избегайте написания всего текста в верхнем регистре")

    return errors


def check_paragraphs_structure(document) -> List[str]:
    # Проверяет структуру абзацев по ГОСТу
    errors = []

    # Проверка длины абзацев (не должны быть слишком длинными)
    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if len(text) > 500:  # Слишком длинный абзац
            errors.append(f"• Стр. {i + 1}: Абзац слишком длинный (разбейте на несколько)")

        # Проверка на отсутствие текста между заголовками
        if i > 0 and _is_heading(paragraph) and _is_heading(document.paragraphs[i - 1]):
            if not document.paragraphs[i - 1].text.strip():
                errors.append(f"• Стр. {i + 1}: Между заголовками должен быть текст")

    return errors


def _is_heading(paragraph) -> bool:
    # Проверяет, является ли абзац заголовком
    return (hasattr(paragraph, 'style') and
            hasattr(paragraph.style, 'name') and
            paragraph.style.name.startswith('Heading'))


def check_lists_formatting(document) -> List[str]:
    # Проверяет оформление списков по ГОСТу
    errors = []

    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()

        # Проверка маркированных списков
        if text.startswith(('•', '-', '—', '–')):
            if not text[1:].strip():  # Пустой элемент списка
                errors.append(f"• Стр. {i + 1}: Пустой элемент списка")

        # Проверка нумерованных списков
        if re.match(r'^\d+[\.\)]', text):
            if not text[2:].strip():  # Пустой элемент списка
                errors.append(f"• Стр. {i + 1}: Пустой элемент нумерованного списка")

    return errors

def check_required_fields(document, doc_type: str) -> List[str]:
    # Проверяет наличие обязательных реквизитов по ГОСТ Р 7.0.97-2016
    errors = []
    full_text = " ".join(p.text for p in document.paragraphs).lower()

    # Обязательные реквизиты для разных типов документов
    REQUIRED_FIELDS = {
        "приказ": ["приказ"],
        "служебная записка": ["служебная записка"],
        "отчёт": ["отчёт", "реферат", "список использованных источников","заключение"]
    }

    required = REQUIRED_FIELDS.get(doc_type, [])
    for field in required:
        if field not in full_text:
            errors.append(f"• Отсутствует обязательный реквизит: '{field}'")

    return errors


def extract_headings(document) -> List[Heading]:
    # Извлекает заголовки из документа на основе стилей.
    # Строго по ГОСТу - точка после номера НЕ допускается.
    headings = []

    for i, paragraph in enumerate(document.paragraphs):
        if not paragraph.text.strip():
            continue

        # Определяем уровень заголовка по стилю
        level = None
        if paragraph.style.name.startswith('Heading'):
            try:
                level = int(paragraph.style.name.split()[-1])
            except (ValueError, IndexError):
                continue

        if level is not None:
            text = paragraph.text.strip()

            # Строгий парсинг по ГОСТу - номер и текст разделяются пробелом
            # Формат: "1 Текст" или "1.1 Текст" (без точки после номера)
            parts = text.split(' ', 1)
            if len(parts) > 1:
                potential_number = parts[0]

                # Проверяем, что это номер в правильном формате (без точки в конце)
                if _is_valid_number_format(potential_number, level):
                    number = potential_number
                    remaining_text = parts[1]
                else:
                    # Если номер с точкой в конце - это ошибка
                    number = ""
                    remaining_text = text
            else:
                # Нет пробела после номера или номер отсутствует
                number = ""
                remaining_text = text

            headings.append(Heading(
                level=level,
                text=remaining_text.strip(),
                paragraph_index=i,
                number=number
            ))

    return headings

def _is_valid_number_format(number: str, level: int) -> bool:
    """
    Проверяет корректность формата номера по ГОСТ.
    Согласно ГОСТ Р 7.0.97-2016:
    """
    # Убираем возможные пробелы в начале/конце
    number = number.strip()

    parts = number.split('.')

    # Уровень должен соответствовать количеству частей
    if len(parts) != level:
        return False

    # Все части должны быть числами
    try:
        for part in parts:
            int(part)
    except ValueError:
        return False

    return True


def _check_sequence(heading: Heading, current_numbers: dict, index: int, all_headings: List[Heading]) -> Optional[str]:
    # Упрощенная и надежная проверка последовательности.
    parts = [int(p) for p in heading.number.split('.')]
    level = heading.level

    # Проверяем только текущий уровень
    current_num = parts[level - 1]

    if level not in current_numbers:
        # Новый уровень - должен начинаться с 1
        if current_num != 1:
            return f"• Стр. {heading.paragraph_index + 1}: Первый номер на уровне {level} должен быть 1, а не {current_num}"
        current_numbers[level] = current_num
    else:
        # Проверяем последовательность
        expected = current_numbers[level] + 1
        if current_num != expected:
            return f"• Стр. {heading.paragraph_index + 1}: Ожидался номер {expected}, а не {current_num} на уровне {level}"
        current_numbers[level] = current_num

    # Сбрасываем более глубокие уровни
    for l in list(current_numbers.keys()):
        if l > level:
            del current_numbers[l]

    return None

def check_numbering(document, doc_type: str = "приказ") -> List[str]:
    # Проверяет правильность нумерации разделов документа.
    errors = []
    headings = extract_headings(document)

    if not headings:
        return ["• Документ не содержит заголовков с нумерацией"]

    # Словарь для отслеживания текущих номеров на каждом уровне
    current_numbers = {}

    for i, heading in enumerate(headings):
        level = heading.level

        # Проверяем формат номера
        if heading.number:
            if not _is_valid_number_format(heading.number, level):
                errors.append(
                    f"• Стр. {heading.paragraph_index + 1}: "
                    f"Неверный формат номера '{heading.number}' для уровня {level}"
                )
                continue

            # Проверяем последовательность
            error_msg = _check_sequence(heading, current_numbers, i, headings)
            if error_msg:
                errors.append(error_msg)
        else:
            errors.append(
                f"• Стр. {heading.paragraph_index + 1}: "
                f"Заголовок уровня {level} не содержит номера"
            )

    return errors